"""DOCX and PDF export for transcriptions, summaries, and mind maps."""

import logging
import re
from typing import Optional

from gigaam_transcriber.data_models import TranscriptionResult, _format_time_txt

logger = logging.getLogger(__name__)

_PDF_CSS = """
@page {
    size: A4;
    margin: 2cm;
}
body {
    font-family: "DejaVu Sans", "Liberation Sans", Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #333;
}
h1 { font-size: 18pt; color: #1a1a1a; }
h2 { font-size: 14pt; color: #333; }
h3 { font-size: 12pt; color: #555; }
table {
    width: 100%;
    border-collapse: collapse;
    margin: 1em 0;
}
th, td {
    border: 1px solid #ddd;
    padding: 6px 10px;
    text-align: left;
    vertical-align: top;
}
th {
    background-color: #f5f5f5;
    font-weight: bold;
}
.speaker {
    font-weight: bold;
    color: #1f77b4;
}
.meta {
    color: #666;
    font-size: 0.9em;
    margin-bottom: 1em;
}
"""


def _transcription_to_pdf_html(result: TranscriptionResult) -> str:
    """Generate HTML suitable for PDF from TranscriptionResult."""
    import html as _html

    has_speakers = any(seg.speaker for seg in result.segments)

    rows = []
    for seg in result.segments:
        start_str = _format_time_txt(seg.start)
        end_str = _format_time_txt(seg.end)
        time_cell = f"{start_str} - {end_str}"
        speaker_cell = _html.escape(seg.speaker) if seg.speaker else ""
        text_cell = _html.escape(seg.text)
        rows.append(
            f"<tr><td>{time_cell}</td>"
            + (f"<td class='speaker'>{speaker_cell}</td>" if has_speakers else "")
            + f"<td>{text_cell}</td></tr>"
        )

    header = (
        "<tr><th>Время</th>"
        + ("<th>Спикер</th>" if has_speakers else "")
        + "<th>Текст</th></tr>"
    )

    meta_parts = [
        f"Длительность: {result.duration:.1f}s",
        f"Модель: {_html.escape(result.model_name)}",
        f"Язык: {_html.escape(result.language)}",
    ]
    if result.metadata.get("source"):
        meta_parts.append(f"Файл: {_html.escape(str(result.metadata['source']))}")

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{_PDF_CSS}</style></head>
<body>
<h1>Транскрипция</h1>
<p class="meta">{' | '.join(meta_parts)}</p>
<table>{header}
{chr(10).join(rows)}
</table>
</body></html>"""


def _md_to_docx_runs(paragraph, text: str):
    """Parse inline markdown (bold **text**) and add runs to paragraph."""
    pattern = r"(\*\*(.+?)\*\*)"
    parts = re.split(pattern, text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            continue
        if i > 0 and parts[i - 1].startswith("**") and parts[i - 1].endswith("**"):
            run = paragraph.add_run(part)
            run.bold = True
        else:
            paragraph.add_run(part)


def _parse_md_to_docx(doc, md_text: str, title: str = "Summary"):
    """Parse markdown text and add to python-docx Document."""
    from docx.shared import Pt

    doc.add_heading(title, level=1)

    lines = md_text.split("\n")
    for line in lines:
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _md_to_docx_runs(para, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            para = doc.add_paragraph(style="List Number")
            _md_to_docx_runs(para, re.sub(r"^\d+\.\s", "", line))
        elif line.strip() == "":
            pass
        else:
            para = doc.add_paragraph()
            _md_to_docx_runs(para, line)


def export_docx_transcription(result: TranscriptionResult, output_path: str) -> str:
    """Export transcription to DOCX format.

    Speaker names are pre-substituted at the _result_from_dict level in routers/exports.py.

    Args:
        result: TranscriptionResult to export
        output_path: Path to save the DOCX file

    Returns:
        Path to the created file
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    doc.add_heading("Транскрипция", level=1)

    meta_para = doc.add_paragraph()
    run = meta_para.add_run(f"Длительность: {result.duration:.1f}s | ")
    run.font.size = Pt(9)
    run = meta_para.add_run(f"Модель: {result.model_name} | ")
    run.font.size = Pt(9)
    run = meta_para.add_run(f"Язык: {result.language}")
    run.font.size = Pt(9)
    if result.metadata.get("source"):
        run = meta_para.add_run(f"\nФайл: {result.metadata['source']}")
        run.font.size = Pt(9)

    has_speakers = any(seg.speaker for seg in result.segments)

    cols = 3 if has_speakers else 2
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Время"
    if has_speakers:
        hdr[1].text = "Спикер"
        hdr[2].text = "Текст"
    else:
        hdr[1].text = "Текст"

    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for seg in result.segments:
        row_cells = table.add_row().cells
        start_str = _format_time_txt(seg.start)
        end_str = _format_time_txt(seg.end)
        row_cells[0].text = f"{start_str} - {end_str}"

        if has_speakers:
            if seg.speaker:
                row_cells[1].text = seg.speaker
                for p in row_cells[1].paragraphs:
                    for r in p.runs:
                        r.bold = True
            row_cells[2].text = seg.text
        else:
            row_cells[1].text = seg.text

    doc.save(output_path)
    return output_path


def export_docx_summary(summary_md: str, title: str, output_path: str) -> str:
    """Export summary markdown to DOCX format.

    Args:
        summary_md: Markdown text of the summary
        title: Document title
        output_path: Path to save the DOCX file

    Returns:
        Path to the created file
    """
    from docx import Document

    doc = Document()
    _parse_md_to_docx(doc, summary_md, title)
    doc.save(output_path)
    return output_path


def export_pdf_transcription(result: TranscriptionResult, output_path: str) -> str:
    """Export transcription to PDF format.

    Uses HTML -> weasyprint pipeline.
    Speaker names are pre-substituted at the _result_from_dict level in routers/exports.py.

    Args:
        result: TranscriptionResult to export
        output_path: Path to save the PDF file

    Returns:
        Path to the created file
    """
    from weasyprint import HTML

    html_content = _transcription_to_pdf_html(result)
    HTML(string=html_content).write_pdf(output_path)
    return output_path


def export_pdf_summary(summary_md: str, title: str, output_path: str) -> str:
    """Export summary markdown to PDF format.

    Args:
        summary_md: Markdown text of the summary
        title: Document title
        output_path: Path to save the PDF file

    Returns:
        Path to the created file
    """
    import markdown as md_lib
    from weasyprint import HTML

    body_html = md_lib.markdown(
        summary_md,
        extensions=["tables", "fenced_code", "nl2br"],
    )

    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{_PDF_CSS}</style></head>
<body>
<h1>{title}</h1>
{body_html}
</body></html>"""

    HTML(string=full_html).write_pdf(output_path)
    return output_path


def export_pdf_mindmap(mindmap_md: str, output_path: str) -> str:
    """Export mind map to PDF format.

    Uses static fallback HTML (no JS) for weasyprint compatibility.

    Args:
        mindmap_md: Markdown text of the mind map
        output_path: Path to save the PDF file

    Returns:
        Path to the created file
    """
    from gigaam_transcriber.mindmap import render_mindmap_fallback
    from weasyprint import HTML

    static_html = render_mindmap_fallback(mindmap_md)

    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{_PDF_CSS}
.mindmap-tree {{ font-family: monospace; white-space: pre-wrap; }}
.mindmap-tree ul {{ list-style: none; padding-left: 1.5em; }}
.mindmap-tree li {{ margin: 0.2em 0; }}
</style></head>
<body>
{static_html}
</body></html>"""

    HTML(string=full_html).write_pdf(output_path)
    return output_path


def export_docx_insights(
    action_items: list[dict],
    decisions: list[dict],
    suggested_steps: list[dict],
    output_path: str,
) -> str:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("Задачи (Action Items)", level=1)
    if action_items:
        table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
        hdr = table.rows[0].cells
        hdr[0].text = "Задача"
        hdr[1].text = "Приоритет"
        hdr[2].text = "Ответственный"
        hdr[3].text = "Срок"
        for item in action_items:
            row = table.add_row().cells
            row[0].text = item.get("task", "")
            row[1].text = item.get("priority", "medium")
            row[2].text = item.get("assignee") or "—"
            row[3].text = item.get("deadline") or "—"
    else:
        doc.add_paragraph("Задачи не найдены.")

    doc.add_heading("Решения (Decisions)", level=1)
    if decisions:
        for d in decisions:
            p = doc.add_paragraph()
            p.add_run(d.get("decision", "")).bold = True
            ctx = d.get("context", "")
            if ctx:
                doc.add_paragraph(f"Контекст: {ctx}", style="List Bullet")
    else:
        doc.add_paragraph("Решения не найдены.")

    doc.add_heading("Рекомендуемые шаги (Suggested Steps)", level=1)
    if suggested_steps:
        category_labels = {
            "followup": "Фоллоу-ап",
            "research": "Исследование",
            "communication": "Коммуникация",
            "planning": "Планирование",
        }
        for i, step in enumerate(suggested_steps, 1):
            p = doc.add_paragraph(style="List Number")
            cat = step.get("category", "planning")
            p.add_run(f"[{category_labels.get(cat, cat)}] ")
            p.add_run(step.get("step", "")).bold = True
            reason = step.get("reason", "")
            if reason:
                doc.add_paragraph(f"Почему: {reason}", style="List Bullet 2")
    else:
        doc.add_paragraph("Рекомендации не найдены.")

    doc.save(output_path)
    return output_path
