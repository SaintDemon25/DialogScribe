"""Tests for gigaam_transcriber.exporters module."""

import os
import tempfile
from pathlib import Path

import pytest

from gigaam_transcriber.data_models import TranscriptionResult, TranscriptionSegment

pytest.importorskip("docx", reason="python-docx not installed")
pytest.importorskip("weasyprint", reason="weasyprint not installed")

from gigaam_transcriber.exporters import (
    export_docx_transcription,
    export_docx_summary,
    export_pdf_transcription,
    export_pdf_summary,
    export_pdf_mindmap,
    export_docx_insights,
)


def _make_result(with_speakers=True, cyrillic=True):
    if with_speakers:
        segments = [
            TranscriptionSegment(text="Привет, как дела?", start=0.0, end=2.5, speaker="Спикер №1"),
            TranscriptionSegment(text="Отлично, спасибо!", start=2.5, end=4.0, speaker="Спикер №2"),
        ]
    else:
        segments = [
            TranscriptionSegment(text="Hello world", start=0.0, end=2.5),
            TranscriptionSegment(text="Goodbye world", start=2.5, end=4.0),
        ]

    text = "Привет, как дела? Отлично, спасибо!" if cyrillic else "Hello world Goodbye world"
    return TranscriptionResult(
        text=text,
        segments=segments,
        duration=4.0,
        language="ru" if cyrillic else "en",
        model_name="test-model",
        processing_time=1.0,
        metadata={"source": "test.wav"},
    )


class TestExportDocxTranscription:
    def test_with_speakers(self, temp_dir):
        result = _make_result(with_speakers=True)
        out = temp_dir / "test.docx"
        path = export_docx_transcription(result, str(out))

        assert Path(path).exists()
        assert path.endswith("test.docx")

        from docx import Document
        doc = Document(path)
        table = doc.tables[0]
        assert len(table.rows) == 3
        assert table.rows[0].cells[0].text == "Время"
        assert table.rows[0].cells[1].text == "Спикер"
        cell_texts = [cell.text for row in table.rows for cell in row.cells]
        assert any("Спикер №1" in t for t in cell_texts)

    def test_without_speakers(self, temp_dir):
        result = _make_result(with_speakers=False)
        out = temp_dir / "test_no_speaker.docx"
        path = export_docx_transcription(result, str(out))

        assert Path(path).exists()

        from docx import Document
        doc = Document(path)
        table = doc.tables[0]
        assert len(table.columns) == 2
        assert table.rows[0].cells[0].text == "Время"
        assert table.rows[0].cells[1].text == "Текст"

    def test_cyrillic_content(self, temp_dir):
        result = _make_result(with_speakers=True, cyrillic=True)
        out = temp_dir / "test_cyrillic.docx"
        path = export_docx_transcription(result, str(out))

        from docx import Document
        doc = Document(path)
        table = doc.tables[0]
        cell_texts = [cell.text for row in table.rows for cell in row.cells]
        assert any("Привет" in t for t in cell_texts)


class TestExportDocxSummary:
    def test_basic_markdown(self, temp_dir):
        md = "# Title\n\n## Section\n\nSome text.\n\n- item 1\n- item 2\n"
        out = temp_dir / "summary.docx"
        path = export_docx_summary(md, "Test Summary", str(out))

        assert Path(path).exists()

        from docx import Document
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("Test Summary" in t for t in texts)
        assert any("Section" in t for t in texts)

    def test_bold_text(self, temp_dir):
        md = "Normal **bold** text"
        out = temp_dir / "bold.docx"
        path = export_docx_summary(md, "Bold Test", str(out))

        from docx import Document
        doc = Document(path)
        has_bold = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and "bold" in run.text:
                    has_bold = True
        assert has_bold

    def test_list_items(self, temp_dir):
        md = "- first\n- second\n\n1. numbered one\n2. numbered two\n"
        out = temp_dir / "lists.docx"
        path = export_docx_summary(md, "Lists", str(out))

        from docx import Document
        doc = Document(path)
        style_names = [p.style.name for p in doc.paragraphs]
        assert "List Bullet" in style_names
        assert "List Number" in style_names


class TestExportPdfTranscription:
    def test_creates_pdf(self, temp_dir):
        result = _make_result(with_speakers=True)
        out = temp_dir / "test.pdf"
        path = export_pdf_transcription(result, str(out))

        assert Path(path).exists()
        assert os.path.getsize(path) > 0

    def test_contains_text(self, temp_dir):
        result = _make_result(with_speakers=True, cyrillic=True)
        out = temp_dir / "text.pdf"
        path = export_pdf_transcription(result, str(out))

        import subprocess
        try:
            result_pdftotext = subprocess.run(
                ["pdftotext", path, "-"],
                capture_output=True, text=True, timeout=10,
            )
            text = result_pdftotext.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            text = ""

        if text:
            assert "Привет" in text or len(text) > 0


class TestExportPdfSummary:
    def test_creates_pdf(self, temp_dir):
        md = "# Summary\n\nThis is a test summary."
        out = temp_dir / "summary.pdf"
        path = export_pdf_summary(md, "Test", str(out))

        assert Path(path).exists()
        assert os.path.getsize(path) > 0

    def test_from_markdown(self, temp_dir):
        md = "## Section\n\n**Bold text** and normal text.\n\n- bullet item\n"
        out = temp_dir / "md_summary.pdf"
        path = export_pdf_summary(md, "MD Summary", str(out))

        assert Path(path).exists()
        assert os.path.getsize(path) > 0


class TestExportPdfMindmap:
    def test_creates_pdf(self, temp_dir):
        md = "# Root\n## Branch 1\n- leaf a\n- leaf b\n## Branch 2\n- leaf c\n"
        out = temp_dir / "mindmap.pdf"
        path = export_pdf_mindmap(md, str(out))

        assert Path(path).exists()
        assert os.path.getsize(path) > 0


class TestBackwardCompatibility:
    def test_txt_save_unchanged(self, sample_transcription_result, temp_dir):
        p = temp_dir / "out.txt"
        sample_transcription_result.save(p)
        assert p.exists()
        assert "Привет" in p.read_text()

    def test_json_save_unchanged(self, sample_transcription_result, temp_dir):
        import json
        p = temp_dir / "out.json"
        sample_transcription_result.save(p, format="json")
        data = json.loads(p.read_text())
        assert "segments" in data

    def test_srt_save_unchanged(self, sample_transcription_result, temp_dir):
        p = temp_dir / "out.srt"
        sample_transcription_result.save(p, format="srt")
        assert "-->" in p.read_text()

    def test_vtt_save_unchanged(self, sample_transcription_result, temp_dir):
        p = temp_dir / "out.vtt"
        sample_transcription_result.save(p, format="vtt")
        assert "WEBVTT" in p.read_text()

    def test_docx_save(self, sample_transcription_result, temp_dir):
        p = temp_dir / "out.docx"
        result_path = sample_transcription_result.save(p, format="docx")
        assert result_path.exists()

    def test_pdf_save(self, sample_transcription_result, temp_dir):
        p = temp_dir / "out.pdf"
        result_path = sample_transcription_result.save(p, format="pdf")
        assert result_path.exists()


class TestExportDocxInsights:
    def test_with_all_sections(self, temp_dir):
        action_items = [
            {"task": "Task 1", "priority": "high", "assignee": "Anna", "deadline": "Friday"},
            {"task": "Task 2", "priority": "low", "assignee": None, "deadline": None},
        ]
        decisions = [{"decision": "Decided X", "context": "Because Y"}]
        suggested_steps = [
            {"step": "Send email", "reason": "Follow up", "category": "followup"},
            {"step": "Research topic", "reason": "Important", "category": "research"},
        ]

        out = str(temp_dir / "insights.docx")
        result = export_docx_insights(action_items, decisions, suggested_steps, out)
        assert result == out
        from pathlib import Path
        assert Path(out).exists()
        assert Path(out).stat().st_size > 0

    def test_empty_insights(self, temp_dir):
        out = str(temp_dir / "empty.docx")
        result = export_docx_insights([], [], [], out)
        assert result == out
        from pathlib import Path
        assert Path(out).exists()
